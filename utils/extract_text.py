import docx
import subprocess
import fitz  # PyMuPDF
import logging
import os

def extraer_texto_docx(archivo):
    doc = docx.Document(archivo)
    textos = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                textos.append(cell.text)

    return "\n".join(textos)

def extraer_texto_rtf(archivo):
    from striprtf.striprtf import rtf_to_text
    with open(archivo, "r", encoding="utf-8") as f:
        rtf_content = f.read()
    return rtf_to_text(rtf_content).strip()

def extraer_texto_pdf(archivo):
    doc = fitz.open(archivo)
    texto = ""
    for pagina in doc:
        texto += pagina.get_text()
    return texto

def extraer_texto_doc(archivo):
    try:
        resultado = subprocess.run(["catdoc", archivo], capture_output=True, text=True)
        if resultado.returncode == 0:
            return resultado.stdout.strip()
        else:
            raise RuntimeError(f"catdoc no pudo procesar {archivo}: {resultado.stderr}")
    except FileNotFoundError:
        logging.error("❌ No se pudo procesar .doc porque 'catdoc' no está instalado.")
        return None
    except Exception as e:
        logging.error(f"❌ Error al procesar .doc {archivo}: {e}")
        return None

def leer_archivo(file_path):
    """Lee el contenido de un archivo y retorna el texto extraído o None si hay un error."""
    
    _, extension = os.path.splitext(file_path)
    texto = ""

    try:
        extension = extension.lower()

        if extension == ".docx":
            texto = extraer_texto_docx(file_path)
        elif extension == ".pdf":
            texto = extraer_texto_pdf(file_path)
        elif extension == ".rtf":
            texto = extraer_texto_rtf(file_path)
        elif extension == ".doc":
            texto = extraer_texto_doc(file_path)
        else:
            logging.error(f"⚠️ Extensión de archivo no soportada: {file_path}")
            return None

        return texto if texto.strip() else None  # Evitar devolver strings vacíos

    except Exception as e:
        logging.error(f"❌ Error inesperado al leer {file_path}: {e}")
        return None