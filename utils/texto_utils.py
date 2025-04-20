from striprtf.striprtf import rtf_to_text
import docx
import subprocess
import logging
import re

def extraer_regex(texto, patron, grupo=1):

    match = re.search(patron, texto, re.IGNORECASE)
    if match:
        return match.group(grupo).strip()
    return None

def normalizar_texto(texto):

    # Diccionario de reemplazos de caracteres especiales por sus equivalentes Unicode
    REEMPLAZOS = {
        'á': 'a',
        'é': 'e',
        'í': 'i',
        'ó': 'o',
        'ú': 'u',
        'Á': 'A',
        'É': 'E',
        'Í': 'I',
        'Ó': 'O',
        'Ú': 'U',
        'ñ': 'n',
        'Ñ': 'N'
    }
    
    # Reemplazar caracteres específicos con equivalentes ASCII
    for original, reemplazo in REEMPLAZOS.items():
        texto = texto.replace(original, reemplazo)

    texto = re.sub(r'\s+', ' ', texto)  # Reemplazar múltiples espacios por uno solo
    texto = re.sub(r'\n+', '|', texto)  # Reemplazar múltiples saltos de línea por uno solo
    texto = re.sub(r'\r+', '|', texto)  # Eliminar retornos de carro
    texto = re.sub(r'\t+', '|', texto)  # Reemplazar múltiples tabulaciones por un espacio
    texto = texto.strip()  # Eliminar espacios en blanco al inicio y al final

    return texto

# ########################################################################### DOCX
def extraer_texto_docx(archivo):
    try:
        doc = docx.Document(archivo)
        textos = []

        # Extraer texto de las tablas
        for table in doc.tables:
            for row in table.rows:
                fila = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if fila:
                    textos.append(fila)

        # Extraer texto de los párrafos
        for paragraph in doc.paragraphs:
            textos.append(paragraph.text)
        
        # Extraer texto del header
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                textos.append(paragraph.text)

        # Extraer texto del footer
        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                textos.append(paragraph.text)
        return "\n".join(textos)

    except Exception as e:
        logging.error(f"Error al procesar el archivo {archivo}: {e}")
    
    return None

# ########################################################################### RTF
def extraer_texto_rtf(archivo):
    try:
        with open(archivo, "r", encoding="utf-8") as f:
            rtf_content = f.read()
        return rtf_to_text(rtf_content).strip()
    except Exception as e:
        logging.error(f"Error al procesar el archivo {archivo}: {e}")
    return None

# ########################################################################### DOC
def extraer_texto_doc(archivo):
    try:
        resultado = subprocess.run(["catdoc", archivo], capture_output=True, text=True)
        if resultado.returncode == 0:
            return resultado.stdout.strip()
        else:
            raise RuntimeError(f"catdoc no pudo procesar {archivo}: {resultado.stderr}")
    except Exception as e:
        logging.error(f"Error al procesar .doc {archivo} $$ {e}")
        return None

# ########################################################################### SUBCADENAS
import re
import logging

def extraer_subcadenas(texto, inicio, fin):
    logging.debug(f"Texto original: {texto}")
    logging.debug(f"Inicio: {inicio}, Fin: {fin}")
    try:
        # Buscar la primera aparición de la cadena de inicio
        match_inicio = re.search(inicio, texto, re.DOTALL | re.IGNORECASE)
        if not match_inicio:
            logging.warning(f"No se encontró la cadena de inicio ({inicio})")
            return []  # No se encontró la cadena de inicio
        logging.debug(f"Cadena de inicio encontrada: {match_inicio.group(0)}")
        # Buscar la primera aparición de la cadena de fin a partir de la posición de la cadena de inicio
        match_fin = re.search(fin, texto[match_inicio.end():], re.DOTALL | re.IGNORECASE)
        if not match_fin:
            logging.warning(f"No se encontró la cadena de fin ({fin}) después de {inicio}")
            return []  # No se encontró la cadena de fin
        logging.debug(f"Cadena de fin encontrada: {match_fin.group(0)}")
        # Extraer la subcadena entre las cadenas de inicio y fin
        subcadena = texto[match_inicio.end():match_inicio.end() + match_fin.start()].strip()

        return subcadena

    except Exception as e:
        logging.error(f"Error al extraer subcadena: {e}")
    
    return None


# ########################################################################### TIPOS DE EXAMEN

def determinar_tipos_examenes(texto: str) -> list:
    try:
        # Cadenas a buscar para determinar tipo de examen
        tipos_examen = {
            "BASAL": [r"INFORME\s+DE\s+POLISOMNOGRAFIA\s+BASAL"],
            "CPAP": [r"INFORME\s+DE\s+POLISOMNOGRAFIA\s+EN\s+TITULACION\s+DE\s+CPAP"],
            "DAM": [r"INFORME\s+DE\s+POLISOMNOGRAFIA\s+BASAL\s+CON\s+DISPOSITIVO(\s+DE\s+AVANCE)?\s+MANDIBULAR"],
            "BPAP": [r"INFORME\s+DE\s+POLISOMNOGRAFIA\s+EN\s+TITULACION\s+DE\s+B[I]?PAP"],
            "ACTIGRAFIA": [r"ESTADISTICAS\s+DIARIAS"],
            "CAPNOGRAFIA": [r"INFORME\s+DE\s+CAPNOGRAFIA"],
            "AUTOCPAP": [r"INFORME\s+DE\s+TITULACION\s+CON\s+AUTO\s+CPAP"],
            "POLIGRAFIA": [r"INFORME\s+POLIGRAFIA\s+RESPIRATORIA"]
        }
        tipos_detectados = []
        for tipo, patrones in tipos_examen.items():
            for patron in patrones:
                if re.search(patron, texto, re.IGNORECASE):
                    logging.debug(f"Tipo de examen {tipo} detectado en el texto {texto}")
                    tipos_detectados.append(tipo)
                    break
        if not tipos_detectados:
            tipos_detectados.append("DESCONOCIDO")
        return tipos_detectados

    except Exception as e:
        logging.error(f"Error al determinar los tipos de examen en el texto: {e}")
    
    return ["ERROR"]
